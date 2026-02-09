import requests
import json
import datetime
import os
from collections import defaultdict
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: 'python-docx' library not found. Please run: pip install python-docx requests")
    exit(1)

# Configuration
API_URL = "http://110.40.129.184/api/records"
# Fallback to local file if API fails (assuming running on same machine as backend dev)
LOCAL_FILE = os.path.join(os.path.dirname(__file__), "backend", "pwa_records.json")
OUTPUT_DIR = os.path.join(os.path.expanduser("~"), "Desktop")

# Targets (Weekly)
TARGETS = {
    'peidong': {'roadshow': 4, 'service': 5}, # Approx 15/4 weeks ~ 4, 20/4 ~ 5
    'xuya': {'roadshow': 4, 'service': 5},
    'xiaoxi': {'roadshow': 4, 'service': 5},
    'tianran': {'roadshow': 4, 'service': 5}
}
# Actually user set Quarterly Targets (15 Roadshows, 20 Services).
# Weekly target approx = Quarterly / 12 weeks?
# User asked to compare "Current Workload" vs "Last Week".
# I will show counts primarily.

def get_date_range(base_date):
    """
    Get Monday and Saturday of the week for the given date.
    Returns (monday_obj, saturday_obj)
    """
    # Create a date object
    if isinstance(base_date, str):
        base_date = datetime.datetime.strptime(base_date, "%Y-%m-%d").date()
    
    # Python weekday: Mon=0, Sun=6
    weekday = base_date.weekday()
    
    monday = base_date - datetime.timedelta(days=weekday)
    saturday = monday + datetime.timedelta(days=5)
    
    return monday, saturday

def fetch_data():
    try:
        print(f"Fetching data from {API_URL}...")
        resp = requests.get(API_URL, timeout=5)
        if resp.status_code == 200:
            return resp.json()
    except Exception as e:
        print(f"API fetch failed: {e}. Trying local file...")
    
    if os.path.exists(LOCAL_FILE):
        with open(LOCAL_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def filter_data_by_range(data, start_date, end_date):
    """
    Filter records between start_date (inclusive) and end_date (inclusive).
    """
    filtered = []
    # Convert range to timestamp or string comp
    # Logic: record['date'] is string "YYYY-MM-DD"
    s_str = start_date.strftime("%Y-%m-%d")
    e_str = end_date.strftime("%Y-%m-%d")
    
    for r in data:
        r_date = r.get('date', '')
        if s_str <= r_date <= e_str:
            filtered.append(r)
    return filtered

def analyze_workload(records):
    # Analyze by member and type
    stats = defaultdict(lambda: defaultdict(int))
    details = defaultdict(list)
    
    for r in records:
        member = r.get('member', 'Unknown')
        rtype = r.get('type', 'other')
        topic = r.get('topic', '')
        inst = r.get('institution', '')
        
        stats[member][rtype] += 1
        stats['Team'][rtype] += 1
        
        # Store details for report
        details[member].append(f"[{rtype}] {inst + ' ' if inst else ''}{topic}")
        
    return stats, details

def generate_word_report(current_week_stats, last_week_stats, current_details, start_date, end_date):
    doc = Document()
    
    # Title
    title = doc.add_heading(level=0)
    run = title.add_run(f"周度数据分析报告 ({start_date.strftime('%m.%d')} - {end_date.strftime('%m.%d')})")
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 153)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Market Overview
    doc.add_heading('一、市场主线复盘', level=1)
    doc.add_paragraph("[此处填写本周市场核心主线及策略观点...]")
    
    # 2. Team Overview
    doc.add_heading('二、团队整体工作量对比', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '本周完成'
    hdr_cells[2].text = '上周完成'
    hdr_cells[3].text = '环比变化'
    
    metrics = [
        ('路演 (Roadshow)', 'roadshow'),
        ('电话会 (Call)', 'call'),
        ('研报 (Report)', 'report'),
        ('高频输出 (Service)', 'service'),
        ('对内服务 (Internal)', 'internal')
    ]
    
    for label, key in metrics:
        row_cells = table.add_row().cells
        curr = current_week_stats['Team'][key]
        last = last_week_stats['Team'][key]
        diff = curr - last
        diff_str = f"{'+' if diff > 0 else ''}{diff}"
        
        row_cells[0].text = label
        row_cells[1].text = str(curr)
        row_cells[2].text = str(last)
        row_cells[3].text = diff_str
    
    # 3. Member Breakdown
    doc.add_heading('三、成员详细工作进度', level=1)
    
    members = ['沛东', '徐亚', '晓希', '天然']
    
    for m in members:
        doc.add_heading(f"{m} 工作周报", level=2)
        
        # Stats Line
        s_curr = current_week_stats[m]
        s_last = last_week_stats[m]
        
        summary = f"本周完成: 路演 {s_curr['roadshow']} (上周{s_last['roadshow']}), " \
                  f"高频 {s_curr['service']} (上周{s_last['service']}), " \
                  f"对内 {s_curr['internal']} (上周{s_last['internal']})"
        doc.add_paragraph(summary)
        
        # Detail List
        if current_details[m]:
            p = doc.add_paragraph()
            p.add_run("本周具体产出：").bold = True
            for item in current_details[m]:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("本周无录入记录。")
            
    # Save
    filename = f"周度数据分析报告_{end_date.strftime('%Y%m%d')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"Report generated successfully: {filepath}")

def main():
    today = datetime.date.today()
    # today = datetime.date(2026, 2, 8) # Debug/Test date
    
    # Calculate This Week (Mon-Sat defined by user)
    # If today is Sunday (weekday=6), "This Week" ended yesterday? 
    # Or "This Week" includes today logic?
    # User said: "Statistics weekly Mon to Sat 23:59".
    # Assuming standard reporting: We report on the FINISHED week.
    # If today is Sunday, we report on the Mon-Sat just passed.
    # If today is Monday, we report on the PREVIOUS Mon-Sat.
    
    weekday = today.weekday()
    if weekday == 6: # Sunday
        # Report for the week ending yesterday
        this_week_monday = today - datetime.timedelta(days=6)
        this_week_saturday = today - datetime.timedelta(days=1)
    elif weekday == 0: # Monday
        # Report for the LAST week
        this_week_monday = today - datetime.timedelta(days=7)
        this_week_saturday = today - datetime.timedelta(days=2)
    else:
        # Report for CURRENT week so far? Or Last week?
        # Usually reports are for completed weeks.
        # But "Statistics weekly... generate a report" implies generating for the current context.
        # Let's assume we report for the week containing TODAY (if Tu-Sa) or Previous (if Su-Mo).
        # Actually user creates report to see "Team Position". Likely Current Week (Real-time).
        # Let's use Current Week Monday to Current Week Saturday (future or passed).
        this_week_monday = today - datetime.timedelta(days=weekday)
        this_week_saturday = this_week_monday + datetime.timedelta(days=5)

    last_week_monday = this_week_monday - datetime.timedelta(days=7)
    last_week_saturday = this_week_saturday - datetime.timedelta(days=7)
    
    print(f"Generating report for period: {this_week_monday} to {this_week_saturday}")
    print(f"Comparing with: {last_week_monday} to {last_week_saturday}")
    
    data = fetch_data()
    if not data:
        print("No data found.")
        return

    current_data = filter_data_by_range(data, this_week_monday, this_week_saturday)
    last_data = filter_data_by_range(data, last_week_monday, last_week_saturday)
    
    curr_stats, curr_details = analyze_workload(current_data)
    last_stats, _ = analyze_workload(last_data)
    
    generate_word_report(curr_stats, last_stats, curr_details, this_week_monday, this_week_saturday)

if __name__ == "__main__":
    main()
